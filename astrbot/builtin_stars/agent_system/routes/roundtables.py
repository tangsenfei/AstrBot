"""
智能体管理模块 - 圆桌会议管理路由

提供圆桌会议管理相关的 REST API
"""
from __future__ import annotations

import json
import asyncio
from datetime import datetime
from typing import TYPE_CHECKING

from quart import request, Response as QuartResponse

from astrbot.core import logger
from astrbot.dashboard.routes.route import Response
from ..models import RoundtableStatus

if TYPE_CHECKING:
    from ..main import AgentSystemPlugin

_plugin_instance: "AgentSystemPlugin | None" = None

# SSE 事件队列管理
_sse_queues: dict[str, asyncio.Queue] = {}
_sse_connections: dict[str, bool] = {}


def register_roundtable_routes(plugin: "AgentSystemPlugin") -> None:
    """注册圆桌会议管理 API 路由

    Args:
        plugin: 插件实例
    """
    global _plugin_instance
    _plugin_instance = plugin

    plugin.context.register_web_api(
        "/agent/roundtables",
        _list_roundtables,
        ["GET"],
        "获取圆桌会议列表"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/<roundtable_id>",
        _get_roundtable,
        ["GET"],
        "获取圆桌会议详情"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/add",
        _create_roundtable,
        ["POST"],
        "创建圆桌会议"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/update",
        _update_roundtable,
        ["POST"],
        "更新圆桌会议"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/delete",
        _delete_roundtable,
        ["POST"],
        "删除圆桌会议"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/<roundtable_id>/execute",
        _execute_roundtable,
        ["POST"],
        "执行圆桌会议"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/<roundtable_id>/stream",
        _stream_roundtable,
        ["GET"],
        "圆桌会议执行流式推送"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/<roundtable_id>/export",
        _export_roundtable,
        ["POST"],
        "导出会议纪要"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/<roundtable_id>/prepare",
        _prepare_roundtable,
        ["POST"],
        "智能材料准备"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/meeting-types",
        _get_meeting_types,
        ["GET"],
        "获取会议类型列表"
    )

    plugin.context.register_web_api(
        "/agent/roundtables/<roundtable_id>/continue",
        _continue_roundtable,
        ["POST"],
        "续会（含验收意见）"
    )

    logger.info("Roundtable management API routes registered")


def _get_roundtable_service():
    """获取 RoundtableService 实例"""
    from ..services.roundtable_service import RoundtableService
    from ..database import get_database

    db = get_database()
    context = None
    if _plugin_instance and _plugin_instance.context:
        context = _plugin_instance.context

    return RoundtableService(db, context)


async def _list_roundtables():
    """获取圆桌会议列表"""
    try:
        service = _get_roundtable_service()
        status = request.args.get("status")
        roundtables = service.get_roundtables(status)
        return Response().ok([r.to_dict() for r in roundtables]).__dict__
    except Exception as e:
        logger.error(f"Failed to list roundtables: {e}")
        return Response().error(str(e)).__dict__


async def _get_roundtable(roundtable_id: str):
    """获取圆桌会议详情

    Args:
        roundtable_id: 圆桌会议 ID
    """
    try:
        service = _get_roundtable_service()
        roundtable = service.get_roundtable(roundtable_id)
        if not roundtable:
            return Response().error(f"圆桌会议 '{roundtable_id}' 不存在").__dict__

        return Response().ok(roundtable.to_dict()).__dict__
    except Exception as e:
        logger.error(f"Failed to get roundtable {roundtable_id}: {e}")
        return Response().error(str(e)).__dict__


async def _create_roundtable():
    """创建圆桌会议

    Request Body:
        {
            "name": "会议名称",
            "topic": "讨论主题",
            "deliverable": "预期产出",
            "mode": "free",
            "host_agent_id": "agent_xxx",
            "participants": ["agent_1", "agent_2"],
            "rounds": 3,
            "config": {...}
        }
    """
    try:
        service = _get_roundtable_service()
        data = await request.get_json()
        if not data:
            return Response().error("请求体不能为空").__dict__

        roundtable = service.create_roundtable(data)
        return Response().ok(roundtable.to_dict(), "圆桌会议创建成功").__dict__
    except ValueError as e:
        return Response().error(str(e)).__dict__
    except Exception as e:
        logger.error(f"Failed to create roundtable: {e}")
        return Response().error(str(e)).__dict__


async def _update_roundtable():
    """更新圆桌会议

    Request Body:
        {
            "id": "会议ID",
            "name": "新名称",
            ...
        }
    """
    try:
        service = _get_roundtable_service()
        data = await request.get_json()
        if not data:
            return Response().error("请求体不能为空").__dict__

        roundtable_id = data.get("id") or data.get("roundtable_id")
        if not roundtable_id:
            return Response().error("缺少圆桌会议 ID").__dict__

        roundtable = service.update_roundtable(roundtable_id, data)
        if not roundtable:
            return Response().error(f"圆桌会议 '{roundtable_id}' 不存在").__dict__

        return Response().ok(roundtable.to_dict(), "圆桌会议更新成功").__dict__
    except ValueError as e:
        return Response().error(str(e)).__dict__
    except Exception as e:
        logger.error(f"Failed to update roundtable: {e}")
        return Response().error(str(e)).__dict__


async def _delete_roundtable():
    """删除圆桌会议

    Request Body:
        {
            "id": "会议ID"
        }
    """
    try:
        service = _get_roundtable_service()
        data = await request.get_json()
        if not data:
            return Response().error("请求体不能为空").__dict__

        roundtable_id = data.get("id") or data.get("roundtable_id")
        if not roundtable_id:
            return Response().error("缺少圆桌会议 ID").__dict__

        success = service.delete_roundtable(roundtable_id)
        if not success:
            return Response().error(f"圆桌会议 '{roundtable_id}' 不存在").__dict__

        return Response().ok(None, "圆桌会议删除成功").__dict__
    except ValueError as e:
        return Response().error(str(e)).__dict__
    except Exception as e:
        logger.error(f"Failed to delete roundtable: {e}")
        return Response().error(str(e)).__dict__


async def _execute_roundtable(roundtable_id: str):
    """执行圆桌会议

    Args:
        roundtable_id: 圆桌会议 ID
    """
    try:
        service = _get_roundtable_service()
        data = await request.get_json() or {}

        # 创建 SSE 队列
        _sse_queues[roundtable_id] = asyncio.Queue()
        _sse_connections[roundtable_id] = True

        # 定义事件回调
        def event_callback(event_type: str, event_data: dict):
            if roundtable_id in _sse_queues:
                try:
                    _sse_queues[roundtable_id].put_nowait({
                        "type": event_type,
                        "data": event_data,
                        "timestamp": asyncio.get_event_loop().time()
                    })
                except Exception:
                    pass

        # 在后台执行圆桌会议
        async def run_execution():
            try:
                result = await service.execute_roundtable(
                    roundtable_id, data, event_callback=event_callback
                )
                # 发送完成事件
                event_callback("completed", result)
            except Exception as e:
                logger.error(f"Roundtable execution error: {e}")
                event_callback("error", {"message": str(e)})
            finally:
                _sse_connections[roundtable_id] = False
                # 等待 SSE 连接关闭
                await asyncio.sleep(5)
                if roundtable_id in _sse_queues:
                    del _sse_queues[roundtable_id]
                if roundtable_id in _sse_connections:
                    del _sse_connections[roundtable_id]

        # 启动后台任务
        asyncio.create_task(run_execution())

        return Response().ok({"message": "执行已启动", "roundtable_id": roundtable_id}).__dict__
    except ValueError as e:
        return Response().error(str(e)).__dict__
    except Exception as e:
        logger.error(f"Failed to execute roundtable {roundtable_id}: {e}")
        return Response().error(str(e)).__dict__


async def _stream_roundtable(roundtable_id: str):
    """圆桌会议执行流式推送 (SSE)

    Args:
        roundtable_id: 圆桌会议 ID
    """
    async def event_generator():
        queue = _sse_queues.get(roundtable_id)
        if not queue:
            # 如果没有队列，返回一个错误事件
            yield f"event: error\ndata: {json.dumps({'message': '没有正在进行的执行'})}\n\n"
            return

        while _sse_connections.get(roundtable_id, False):
            try:
                # 等待事件，超时 30 秒
                event = await asyncio.wait_for(queue.get(), timeout=30.0)
                data = json.dumps(event, ensure_ascii=False)
                yield f"event: message\ndata: {data}\n\n"

                # 如果是完成或错误事件，结束流
                if event["type"] in ("completed", "error"):
                    break
            except asyncio.TimeoutError:
                # 发送心跳保持连接
                yield f"event: heartbeat\ndata: {{}}\n\n"
            except Exception as e:
                logger.error(f"SSE stream error: {e}")
                break

    response = QuartResponse(
        event_generator(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Transfer-Encoding": "chunked",
        }
    )
    response.timeout = None
    return response


async def _export_roundtable(roundtable_id: str):
    """导出会议纪要

    Args:
        roundtable_id: 圆桌会议 ID

    Request Body:
        {
            "type": "summary" | "deliverable",  // 导出类型：会议纪要 或 交付物
            "format": "markdown" | "word"       // 导出格式
        }
    """
    try:
        service = _get_roundtable_service()
        roundtable = service.get_roundtable(roundtable_id)
        if not roundtable:
            return Response().error(f"圆桌会议 '{roundtable_id}' 不存在").__dict__

        data = await request.get_json() or {}
        export_type = data.get("type", "summary")  # summary or deliverable
        export_format = data.get("format", roundtable.export_format or "markdown")

        result = roundtable.result or {}

        if export_type == "deliverable":
            content = result.get("deliverable", "") or result.get("summary", "")
            filename = f"{roundtable.name}_交付物"
        else:
            # 生成会议纪要
            content = _generate_meeting_summary(roundtable)
            filename = f"{roundtable.name}_会议纪要"

        if export_format == "word":
            try:
                from docx import Document
                from docx.shared import Pt
                from io import BytesIO

                doc = Document()
                doc.add_heading(filename, 0)

                # 添加会议基本信息
                doc.add_heading("会议信息", level=1)
                info_table = doc.add_table(rows=1, cols=2)
                info_table.style = "Light Grid Accent 1"
                hdr_cells = info_table.rows[0].cells
                hdr_cells[0].text = "属性"
                hdr_cells[1].text = "值"

                rows = [
                    ("会议名称", roundtable.name),
                    ("会议主题", roundtable.topic),
                    ("会议类型", roundtable.meeting_type),
                    ("预期产出", roundtable.deliverable),
                    ("创建时间", roundtable.created_at.strftime("%Y-%m-%d %H:%M:%S")),
                ]
                for key, val in rows:
                    row_cells = info_table.add_row().cells
                    row_cells[0].text = key
                    row_cells[1].text = str(val)

                doc.add_heading("会议内容", level=1)
                # 将 markdown 内容分段
                for line in content.split("\n"):
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("- "):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    elif line.strip():
                        doc.add_paragraph(line)

                # 保存到内存
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                return QuartResponse(
                    buffer.getvalue(),
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f"attachment; filename={filename}.docx",
                    }
                )
            except ImportError:
                # 如果没有 python-docx，回退到 markdown
                export_format = "markdown"
                content = content + "\n\n> 注意：Word 导出需要安装 python-docx 库"

        # Markdown 格式
        return QuartResponse(
            content,
            mimetype="text/markdown",
            headers={
                "Content-Disposition": f"attachment; filename={filename}.md",
            }
        )
    except Exception as e:
        logger.error(f"Failed to export roundtable {roundtable_id}: {e}")
        return Response().error(str(e)).__dict__


def _generate_meeting_summary(roundtable) -> str:
    """生成会议纪要 Markdown"""
    lines = []
    lines.append(f"# {roundtable.name} - 会议纪要")
    lines.append("")
    lines.append("## 会议信息")
    lines.append(f"- **会议主题**: {roundtable.topic}")
    lines.append(f"- **会议类型**: {roundtable.meeting_type}")
    lines.append(f"- **预期产出**: {roundtable.deliverable}")
    lines.append(f"- **创建时间**: {roundtable.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")

    # 准备阶段
    if roundtable.preparation_records:
        lines.append("## 准备阶段")
        for record in roundtable.preparation_records:
            lines.append(f"### {record.get('question', '')}")
            lines.append(f"{record.get('answer', '')}")
            lines.append("")

    # 进行阶段
    result = roundtable.result or {}
    discussion_rounds = result.get("discussion_rounds", [])
    if discussion_rounds:
        lines.append("## 会议过程")
        for r in discussion_rounds:
            lines.append(f"### 第 {r.get('round', '?')} 轮")
            speaker = r.get("speaker", "")
            content = r.get("content", "")
            if speaker:
                lines.append(f"**{speaker}**: {content}")
            else:
                lines.append(content)
            lines.append("")

    # 完成阶段
    lines.append("## 会议结论")
    if result.get("summary"):
        lines.append(result["summary"])
    lines.append("")

    if result.get("deliverable"):
        lines.append("## 会议交付物")
        lines.append(result["deliverable"])

    return "\n".join(lines)


async def _prepare_roundtable(roundtable_id: str):
    """智能材料准备

    Args:
        roundtable_id: 圆桌会议 ID

    Request Body:
        {
            "action": "generate_questions" | "submit_answer",
            "data": {...}
        }
    """
    try:
        service = _get_roundtable_service()
        roundtable = service.get_roundtable(roundtable_id)
        if not roundtable:
            return Response().error(f"圆桌会议 '{roundtable_id}' 不存在").__dict__

        req_data = await request.get_json() or {}
        action = req_data.get("action", "generate_questions")

        if action == "generate_questions":
            # 生成引导问题
            questions = _generate_preparation_questions(roundtable)
            return Response().ok({"questions": questions}).__dict__

        elif action == "submit_answer":
            # 提交用户回答
            data = req_data.get("data", {})
            question = data.get("question", "")
            answer = data.get("answer", "")

            records = list(roundtable.preparation_records or [])
            records.append({"question": question, "answer": answer, "time": datetime.now().isoformat()})

            service.update_roundtable(roundtable_id, {"preparation_records": records})
            return Response().ok(None, "回答已记录").__dict__

        else:
            return Response().error(f"未知的 action: {action}").__dict__

    except Exception as e:
        logger.error(f"Failed to prepare roundtable {roundtable_id}: {e}")
        return Response().error(str(e)).__dict__


def _generate_preparation_questions(roundtable) -> list[dict]:
    """根据会议类型生成准备问题"""
    meeting_type = roundtable.meeting_type
    topic = roundtable.topic
    deliverable = roundtable.deliverable

    questions_map = {
        "brainstorm": [
            {"id": "target_user", "question": f"为了「{topic}」的头脑风暴更高效，请描述目标用户群体是谁？", "example": "例如：25-35岁的互联网从业者", "required": False},
            {"id": "constraints", "question": "有哪些技术或资源约束需要考虑？", "example": "例如：预算限制在10万以内，开发周期3个月", "required": False},
            {"id": "references", "question": "是否有竞品或参考案例？", "example": "例如：类似产品A、B的主要特点", "required": False},
        ],
        "parliament": [
            {"id": "options", "question": f"关于「{topic}」，有哪些备选方案？", "example": "例如：方案A使用微服务，方案B使用单体架构", "required": False},
            {"id": "criteria", "question": "评估标准是什么？", "example": "例如：性能、成本、维护性、扩展性", "required": False},
            {"id": "research", "question": "有无已做的初步调研？", "example": "例如：已测试过方案A的POC，性能满足要求", "required": False},
        ],
        "convergence": [
            {"id": "current_status", "question": f"关于「{topic}」，当前的状态是什么？", "example": "例如：已完成需求分析，正在技术选型阶段", "required": False},
            {"id": "constraints", "question": "有哪些必须满足的约束条件？", "example": "例如：必须在Q3上线，团队有5名后端工程师", "required": False},
            {"id": "success_criteria", "question": "如何定义方案成功？", "example": "例如：日活用户达到1万，系统可用性99.9%", "required": False},
        ],
        "six_hat": [
            {"id": "background", "question": f"关于「{topic}」，请提供已知的背景信息", "example": "例如：市场现状、相关数据、历史决策", "required": False},
            {"id": "stakeholders", "question": "涉及哪些利益相关方？", "example": "例如：产品团队、技术团队、运营团队", "required": False},
        ],
        "fishbone": [
            {"id": "problem_detail", "question": f"关于「{topic}」，请描述问题发生的具体场景", "example": "例如：每天上午9-10点高峰期，系统响应时间超过5秒", "required": False},
            {"id": "investigated", "question": "已排查过哪些方面？", "example": "例如：已检查数据库慢查询，已优化索引", "required": False},
            {"id": "data", "question": "有无相关日志或监控数据？", "example": "例如：Grafana监控链接，错误日志片段", "required": False},
        ],
        "swot": [
            {"id": "background", "question": f"关于「{topic}」，请提供背景信息", "example": "例如：公司现状、市场地位、竞争对手情况", "required": False},
            {"id": "timeframe", "question": "分析的时间范围是？", "example": "例如：未来1年、未来3年", "required": False},
        ],
        "okr": [
            {"id": "objective_hint", "question": f"关于「{topic}」，期望达成的目标方向是？", "example": "例如：提升用户留存率、拓展新市场", "required": False},
            {"id": "team_size", "question": "团队规模和资源配置？", "example": "例如：10人产品技术团队，预算50万", "required": False},
            {"id": "timeline", "question": "目标达成的时间节点？", "example": "例如：Q4末完成，明年Q1验证效果", "required": False},
        ],
        "retrospective": [
            {"id": "project_info", "question": f"关于「{topic}」，请描述项目/迭代的基本情况", "example": "例如：2个月完成，团队5人，交付了X功能", "required": False},
            {"id": "metrics", "question": "有哪些量化指标？", "example": "例如：按时交付率、bug数量、用户满意度", "required": False},
        ],
        "interview": [
            {"id": "jd", "question": f"关于「{topic}」岗位，请提供岗位描述", "example": "例如：负责核心系统开发，要求5年以上经验", "required": False},
            {"id": "skills", "question": "必备技能和加分项？", "example": "例如：必备：Java、Redis；加分：K8s、Go", "required": False},
            {"id": "tech_stack", "question": "团队技术栈是什么？", "example": "例如：Spring Cloud、MySQL、RocketMQ", "required": False},
            {"id": "salary", "question": "期望薪资范围？", "example": "例如：30-50K", "required": False},
        ],
        "standard": [
            {"id": "background", "question": f"关于「{topic}」，请补充相关背景信息", "example": "例如：决策背景、已知信息、关注重点", "required": False},
        ],
    }

    return questions_map.get(meeting_type, questions_map["standard"])


async def _get_meeting_types():
    """获取会议类型列表"""
    try:
        meeting_types = [
            {"type": "standard", "name": "标准讨论", "description": "多智能体围绕主题进行多轮自由讨论，逐步深入并达成共识"},
            {"type": "brainstorm", "name": "头脑风暴", "description": "鼓励发散思维，快速生成大量创意和想法，不做评判"},
            {"type": "parliament", "name": "议会辩论", "description": "正反双方就议题展开辩论，通过质询和反驳深化论证"},
            {"type": "convergence", "name": "收敛决策", "description": "从多个方案中逐步筛选、评估，最终收敛到最优决策"},
            {"type": "six_hat", "name": "六顶思考帽", "description": "按六种思维角色（白红黑黄绿蓝）依次分析问题"},
            {"type": "fishbone", "name": "鱼骨分析", "description": "从人、机、料、法、环、测六个维度分析问题根因"},
            {"type": "swot", "name": "SWOT 分析", "description": "从优势、劣势、机会、威胁四个维度进行战略分析"},
            {"type": "okr", "name": "OKR 制定", "description": "制定目标与关键成果，确保目标可衡量、可追踪"},
            {"type": "retrospective", "name": "回顾复盘", "description": "回顾过去的工作，总结经验教训，制定改进计划"},
            {"type": "interview", "name": "面试评审", "description": "模拟面试场景，从多角度评估候选人或方案"},
        ]
        return Response().ok(meeting_types).__dict__
    except Exception as e:
        logger.error(f"Failed to get meeting types: {e}")
        return Response().error(str(e)).__dict__


async def _continue_roundtable(roundtable_id: str):
    """续会

    Args:
        roundtable_id: 圆桌会议 ID

    Request Body:
        {
            "review_comment": "验收意见",
            "additional_rounds": 2,  // 额外增加的轮数
            "additional_topic": "",  // 追加讨论主题（可选）
        }
    """
    try:
        service = _get_roundtable_service()
        roundtable = service.get_roundtable(roundtable_id)
        if not roundtable:
            return Response().error(f"圆桌会议 '{roundtable_id}' 不存在").__dict__

        if roundtable.status not in (RoundtableStatus.COMPLETED, RoundtableStatus.FAILED):
            return Response().error("只有已完成或失败的会议才能续会").__dict__

        data = await request.get_json() or {}
        review_comment = data.get("review_comment", "")
        additional_rounds = data.get("additional_rounds", 2)
        additional_topic = data.get("additional_topic", "")

        # 保存历史记录
        previous_result = roundtable.result or {}
        previous_records = roundtable.discussion_records or []

        # 构建续会上下文
        continuation_context = {
            "previous_summary": previous_result.get("summary", ""),
            "previous_deliverable": previous_result.get("deliverable", ""),
            "review_comment": review_comment,
            "continuation_time": datetime.now().isoformat(),
        }

        # 更新会议状态为待执行，保留历史
        update_data = {
            "status": RoundtableStatus.PENDING.value,
            "rounds": roundtable.rounds + additional_rounds,
            "stage": "pending",
            "current_round": 0,
            "current_speaker": "",
            "streaming_content": "",
        }

        # 如果有追加主题，更新 topic
        if additional_topic:
            update_data["topic"] = f"{roundtable.topic}\n\n【续会追加议题】{additional_topic}"

        # 将验收意见和续会信息存入 preparation_records
        prep_records = list(roundtable.preparation_records or [])
        prep_records.append({
            "question": "续会验收意见",
            "answer": review_comment or "无特别意见，继续讨论",
            "time": datetime.now().isoformat(),
            "type": "review",
        })
        if additional_topic:
            prep_records.append({
                "question": "续会追加议题",
                "answer": additional_topic,
                "time": datetime.now().isoformat(),
                "type": "continuation",
            })
        update_data["preparation_records"] = prep_records

        # 保存续会上下文到 config
        config = dict(roundtable.config or {})
        if "continuations" not in config:
            config["continuations"] = []
        config["continuations"].append(continuation_context)
        update_data["config"] = config

        # 保留之前的讨论记录到 config 中
        config["previous_discussion_records"] = previous_records
        config["previous_result"] = previous_result

        service.update_roundtable(roundtable_id, update_data)

        return Response().ok(None, "续会已就绪，可以开始执行").__dict__

    except Exception as e:
        logger.error(f"Failed to continue roundtable {roundtable_id}: {e}")
        return Response().error(str(e)).__dict__
